# generar_acta.py
import os
from pydantic_ai.mcp import mcp_server, tool
from docx import Document
from datetime import datetime
from typing import Optional
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider
from dotenv import load_dotenv

load_dotenv(override=True)

# Inicializar proveedor DeepSeek (mismo que en app.py)
deepseek_provider = OpenAIProvider(
    base_url='https://api.deepseek.com',
    api_key=os.environ["DEEPSEEK_API_KEY"]
)

deepseek_chat_model = OpenAIModel(
    'deepseek-chat',
    provider=deepseek_provider
)

@mcp_server
class ActasServer:
    @tool
    async def generar_acta_desde_transcripcion(
        self,
        transcripcion: str,
        titulo: Optional[str] = "Acta de reunión"
    ) -> str:
        """
        Genera un acta en formato Word a partir de la transcripción completa de la reunión.
        Extrae automáticamente fecha, participantes, temas y acuerdos usando el modelo LLM.
        """

        # Pedimos al LLM que estructure la información
        prompt = f"""
        Eres un asistente que analiza transcripciones de reuniones.
        Con la siguiente transcripción:
        ---
        {transcripcion}
        ---
        Extrae la información en JSON con las claves:
        fecha (YYYY-MM-DD),
        participantes (lista de nombres),
        temas (lista de temas tratados),
        acuerdos (lista de acuerdos alcanzados).
        """
        result = await deepseek_chat_model.run(prompt)

        # Parsear JSON del resultado
        import json
        try:
            datos = json.loads(result.data.strip())
        except Exception:
            raise ValueError("No se pudo interpretar la respuesta del modelo como JSON.")

        fecha = datos.get("fecha", datetime.now().strftime("%Y-%m-%d"))
        participantes = datos.get("participantes", [])
        temas = datos.get("temas", [])
        acuerdos = datos.get("acuerdos", [])

        # Crear el documento Word
        doc = Document()
        doc.add_heading(titulo, 0)
        doc.add_paragraph(f"Fecha: {fecha}")

        doc.add_paragraph("Participantes:")
        for p in participantes:
            doc.add_paragraph(f"- {p}", style="List Bullet")

        doc.add_heading("Temas tratados", level=1)
        for t in temas:
            doc.add_paragraph(f"- {t}")

        doc.add_heading("Acuerdos", level=1)
        for a in acuerdos:
            doc.add_paragraph(f"- {a}")

        filename = f"acta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)

        return f"Acta generada en: {filename}"

if __name__ == "__main__":
    ActasServer().run()
